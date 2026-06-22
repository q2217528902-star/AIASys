"""共享文档提取服务。"""

from __future__ import annotations

import logging
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional, Sequence

from app.utils.path_utils import as_system_path

from app.core.config import (
    DOCUMENT_EXTRACTION_DEFAULT_MODE,
    DOCUMENT_EXTRACTION_FALLBACK_MODES,
    DOCUMENT_EXTRACTION_PDF_PASSWORD,
)

from .models import (
    DocumentExtractionMode,
    DocumentExtractionResult,
    DocumentExtractionSettings,
)

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".log",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".go",
    ".rs",
    ".sh",
    ".sql",
    ".xml",
    ".html",
    ".htm",
    ".rst",
}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx", ".doc"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | SPREADSHEET_EXTENSIONS


class DocumentExtractionService:
    """统一文件提取入口。"""

    SUPPORTED_EXTENSIONS = SUPPORTED_EXTENSIONS

    def __init__(self, settings: Optional[DocumentExtractionSettings] = None) -> None:
        self.settings = settings or DocumentExtractionSettings.from_values(
            default_mode=DOCUMENT_EXTRACTION_DEFAULT_MODE,
            fallback_modes=DOCUMENT_EXTRACTION_FALLBACK_MODES,
            pdf_password=DOCUMENT_EXTRACTION_PDF_PASSWORD,
        )

    def extract(
        self,
        file_path: str | Path,
        file_bytes: Optional[bytes] = None,
        *,
        mode: Optional[str | DocumentExtractionMode] = None,
        fallback_modes: Optional[Sequence[str | DocumentExtractionMode]] = None,
    ) -> DocumentExtractionResult:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {suffix or '[no-extension]'}")

        content = file_bytes if file_bytes is not None else Path(as_system_path(str(path))).read_bytes()
        requested_mode = (
            DocumentExtractionMode.parse(mode) if mode is not None else self.settings.default_mode
        )
        mode_chain = self._resolve_mode_chain(requested_mode, fallback_modes)

        errors: list[str] = []
        for current_mode in mode_chain:
            try:
                text = self._extract_with_mode(path, content, current_mode)
                normalized_text = self._normalize_text(text)
                if not normalized_text:
                    raise ValueError("文档解析后为空内容")

                return DocumentExtractionResult(
                    text=normalized_text,
                    mode_used=current_mode,
                    requested_mode=requested_mode,
                    file_type=suffix.lstrip(".") or "unknown",
                    warnings=self._fallback_warnings(
                        requested_mode=requested_mode,
                        mode_used=current_mode,
                        errors=errors,
                    ),
                )
            except Exception as exc:
                errors.append(f"{current_mode.value}: {exc}")
                logger.warning(
                    "文档提取失败，mode=%s, file=%s, error=%s", current_mode.value, path.name, exc
                )

        detail = " | ".join(errors) if errors else "unknown error"
        raise ValueError(f"文档提取失败: {detail}")

    def _resolve_mode_chain(
        self,
        requested_mode: DocumentExtractionMode,
        fallback_modes: Optional[Sequence[str | DocumentExtractionMode]],
    ) -> tuple[DocumentExtractionMode, ...]:
        chain = [requested_mode]
        configured_fallbacks = (
            self.settings.fallback_modes
            if fallback_modes is None
            else DocumentExtractionMode.parse_many(fallback_modes)
        )
        chain.extend(configured_fallbacks)
        return tuple(dict.fromkeys(chain))

    @staticmethod
    def _fallback_warnings(
        *,
        requested_mode: DocumentExtractionMode,
        mode_used: DocumentExtractionMode,
        errors: list[str],
    ) -> list[str]:
        if not errors or mode_used == requested_mode:
            return []
        return [f"已从 {requested_mode.value} 回退到 {mode_used.value}"]

    def _extract_with_mode(
        self,
        path: Path,
        content: bytes,
        mode: DocumentExtractionMode,
    ) -> str:
        if mode == DocumentExtractionMode.DOCLING:
            return self._extract_with_docling(path, content)
        if mode == DocumentExtractionMode.BASIC:
            return self._extract_basic(path.suffix.lower(), content)
        if mode == DocumentExtractionMode.ENHANCED:
            return self._extract_enhanced(path.suffix.lower(), content)
        raise ValueError(f"未知提取模式: {mode.value}")

    def _extract_basic(self, suffix: str, content: bytes) -> str:
        if suffix in TEXT_EXTENSIONS:
            return self._extract_text(content)
        if suffix in PDF_EXTENSIONS:
            return self._extract_pdf(content)
        if suffix in DOCX_EXTENSIONS:
            return self._extract_docx_basic(content)
        if suffix in SPREADSHEET_EXTENSIONS:
            return self._extract_spreadsheet(content)
        raise ValueError(f"不支持的文件类型: {suffix}")

    def _extract_enhanced(self, suffix: str, content: bytes) -> str:
        if suffix in TEXT_EXTENSIONS:
            return self._extract_text(content)
        if suffix in PDF_EXTENSIONS:
            return self._extract_pdf(content)
        if suffix in DOCX_EXTENSIONS:
            return self._extract_docx_enhanced(content)
        if suffix in SPREADSHEET_EXTENSIONS:
            return self._extract_spreadsheet(content)
        raise ValueError(f"不支持的文件类型: {suffix}")

    @staticmethod
    def _extract_text(content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("无法识别文本编码")

    def _extract_pdf(self, content: bytes) -> str:
        try:
            import pypdf  # type: ignore
        except Exception as exc:
            raise ValueError("解析 PDF 需要安装 pypdf") from exc

        try:
            reader = pypdf.PdfReader(BytesIO(content))
            if reader.is_encrypted:
                password = self.settings.pdf_password
                if not password:
                    raise ValueError("PDF 已加密，且未配置解密密码")
                if reader.decrypt(password) == 0:
                    raise ValueError("PDF 解密失败，请检查密码")

            pages: list[str] = []
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(f"--- Page {index} ---\n{text}")
            return "\n\n".join(pages)
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"PDF 解析失败: {exc}") from exc

    @staticmethod
    def _extract_docx_basic(content: bytes) -> str:
        try:
            import docx  # type: ignore
        except Exception as exc:
            raise ValueError("解析 DOCX 需要安装 python-docx") from exc

        try:
            doc = docx.Document(BytesIO(content))
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                rows = [
                    " | ".join(cell.text.strip() for cell in row.cells).strip()
                    for row in table.rows
                ]
                rows = [row for row in rows if row]
                if rows:
                    paragraphs.append("--- Table ---")
                    paragraphs.extend(rows)
            return "\n\n".join(paragraphs)
        except Exception as exc:
            raise ValueError(f"DOCX 解析失败: {exc}") from exc

    @staticmethod
    def _extract_docx_enhanced(content: bytes) -> str:
        try:
            from docx import Document  # type: ignore
            from docx.table import Table  # type: ignore
            from docx.text.paragraph import Paragraph  # type: ignore
        except Exception as exc:
            raise ValueError("解析 DOCX 需要安装 python-docx") from exc

        def escape_cell(value: Optional[str]) -> str:
            text = (value or "").strip()
            return (
                text.replace("\\", "\\\\")
                .replace("\t", "    ")
                .replace("\r\n", "<br>")
                .replace("\r", "<br>")
                .replace("\n", "<br>")
            )

        try:
            doc = Document(BytesIO(content))
            parts: list[str] = []
            for element in doc.element.body.iterchildren():
                if element.tag.endswith("}p"):
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    if text:
                        parts.append(text)
                    continue

                if element.tag.endswith("}tbl"):
                    table = Table(element, doc)
                    rows: list[str] = []
                    for row in table.rows:
                        cells = [escape_cell(cell.text) for cell in row.cells]
                        if any(cell for cell in cells):
                            rows.append("\t".join(cells).rstrip())
                    if rows:
                        parts.append("--- Table ---")
                        parts.extend(rows)

            return "\n\n".join(part for part in parts if part.strip())
        except Exception as exc:
            raise ValueError(f"DOCX 解析失败: {exc}") from exc

    @staticmethod
    def _extract_spreadsheet(content: bytes) -> str:
        try:
            import openpyxl  # type: ignore
        except Exception as exc:
            raise ValueError("解析 XLSX 需要安装 openpyxl") from exc

        try:
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets:
                lines.append(f"# sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    if not any(cell is not None and str(cell).strip() for cell in row):
                        continue
                    cells = [str(cell).strip() if cell is not None else "" for cell in row]
                    lines.append("\t".join(cells).rstrip())
            return "\n".join(lines)
        except Exception as exc:
            raise ValueError(f"XLSX 解析失败: {exc}") from exc

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [line.rstrip() for line in text.replace("\x00", "").splitlines()]
        normalized = "\n".join(lines).strip()
        return normalized

    @staticmethod
    def _extract_with_docling(path: Path, content: bytes) -> str:
        try:
            from docling.document_converter import DocumentConverter  # type: ignore
        except Exception as exc:
            raise ValueError("DOCLING 模式不可用：未安装 docling") from exc

        temp_path: Optional[Path] = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=path.suffix) as temp_file:
                temp_file.write(content)
                temp_path = Path(temp_file.name)

            converter = DocumentConverter()
            result = converter.convert(temp_path)
            return result.document.export_to_markdown()
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"DOCLING 提取失败: {exc}") from exc
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)


_document_extraction_service: Optional[DocumentExtractionService] = None


def get_document_extraction_service() -> DocumentExtractionService:
    global _document_extraction_service
    if _document_extraction_service is None:
        _document_extraction_service = DocumentExtractionService()
    return _document_extraction_service
