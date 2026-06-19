from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest

from app.document_extraction import (
    DocumentExtractionMode,
    DocumentExtractionService,
    DocumentExtractionSettings,
)
from app.knowledge.parser import DocumentParser


def _build_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph("before table")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "name"
    table.rows[0].cells[1].text = "value"
    table.rows[1].cells[0].text = "foo"
    table.rows[1].cells[1].text = "bar"
    document.add_paragraph("after table")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet.append(["name", "score"])
    sheet.append(["alice", 98])

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def test_docx_modes_preserve_different_table_order() -> None:
    service = DocumentExtractionService()
    content = _build_docx_bytes()

    basic = service.extract("sample.docx", content, mode="basic")
    enhanced = service.extract("sample.docx", content, mode="enhanced")

    assert basic.mode_used == DocumentExtractionMode.BASIC
    assert enhanced.mode_used == DocumentExtractionMode.ENHANCED
    assert basic.text.index("after table") < basic.text.index("--- Table ---")
    assert enhanced.text.index("before table") < enhanced.text.index("--- Table ---")
    assert enhanced.text.index("--- Table ---") < enhanced.text.index("after table")


def test_default_mode_uses_configured_mode_for_spreadsheet() -> None:
    service = DocumentExtractionService(
        DocumentExtractionSettings.from_values(
            default_mode="enhanced",
            fallback_modes=["basic"],
        )
    )

    result = service.extract("report.xlsx", _build_xlsx_bytes())

    assert result.mode_used == DocumentExtractionMode.ENHANCED
    assert "# sheet: Data" in result.text
    assert "alice\t98" in result.text


def test_default_mode_falls_back_to_basic_when_primary_mode_fails(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    service = DocumentExtractionService(
        DocumentExtractionSettings.from_values(
            default_mode="docling",
            fallback_modes=["basic"],
        )
    )

    def fake_docling(path: Path, content: bytes) -> str:
        raise ValueError("docling unavailable")

    monkeypatch.setattr(service, "_extract_with_docling", fake_docling)

    result = service.extract("notes.txt", b"hello world")

    assert result.mode_used == DocumentExtractionMode.BASIC
    assert result.text == "hello world"
    assert result.warnings == ["已从 docling 回退到 basic"]


def test_auto_mode_is_not_accepted() -> None:
    service = DocumentExtractionService()

    with pytest.raises(ValueError, match="不支持的 extraction_mode"):
        service.extract("notes.txt", b"hello world", mode="auto")


def test_document_parser_wrapper_keeps_compatibility() -> None:
    text = DocumentParser.parse(Path("demo.txt"), b"compatibility text", mode="basic")
    assert text == "compatibility text"
